import datetime
from io import BytesIO
import os
import re
import json

from django.core.cache import cache
from django.http import HttpResponse

from bs4 import BeautifulSoup
from Levenshtein import jaro_winkler
from docx import Document
import errno
from requests_toolbelt.multipart.encoder import MultipartEncoder, MultipartEncoderMonitor
from tqdm import tqdm
import urllib3
import requests

from .c.app_config import config_url
from .models import LANGUAGE_CHOICES, Article, ArticleCase


class NewArticle():
    '''Parse all data to Article model by url and for comparison of articles'''

    def __init__(self, url=None, soup=None, compare=False):
        print('init new article: ', url)
        self.url = url
        self.soup = soup or BeautifulSoup(
                                requests.get(url).text,
                                'lxml')
        NewArticle.compare_parse(self) if compare else NewArticle.parse(self)

    def parse(self):
        self.title = self.soup.select('h2.entry-title')[0].get_text().strip()
        self.lead = self.soup.select('p.entry-lead')[0].get_text().strip()
        self.text = self.soup.select('div.entry-content')[0].get_text().strip()
        self.symbols_amount = sum(map(len, [self.text, self.title, self.lead]))
        self.img_url = self.soup.article.select('img.img-responsive')[0].get('src')
        for num, i in LANGUAGE_CHOICES:
            if 'anf'+i in self.url:
                self.language = i

        NewArticle.compare_parse(self)

    def compare_parse(self):
        self.url_title = self.url.split('/')[-1]

    def __str__(self):
        return self.url_title


class Manager(NewArticle):
    '''Check url and manage action with article in order to its status'''

    def __init__(self,
            log='tr_helper/last_article_log.txt',
            other_log = 'tr_helper/other_language_log.txt',
            url=None, soup=None):
        NewArticle.__init__(self, url, soup)
        self.log = log
        self.other_log = other_log

    def is_new(self, user_req=False, test=False):
        '''Compare title with titles of alredy existed articles'''

        with open(self.log) as log, open(self.other_log) as other_log:
            for url in other_log.readlines()+log.readlines()[::-1]:
                comp_url_title = url.split('/')[-1]

                text_index = jaro_winkler(
                    self.url_title,
                    comp_url_title)

                if text_index == 1:
                    if user_req:
                        return url
                    print(url)
                    self.update_url = url.strip()
                    return False

        if user_req:
            with open(self.other_log, 'a') as other_log:
                other_log.write(self.url+'\n')
            return False

        return True

    def update(self):
        print('update')
        query = Article.objects.get(url=self.update_url)
        query.url = self.url
        query.title = self.title
        query.symbols_amount = self.symbols_amount
        query.save()

    def write_bd(self):
        print('write')
        self.case = ArticleCase()
        self.case.save()
        new = Article(
            case = self.case,
            url=self.url,
            title=self.title,
            symbols_amount=self.symbols_amount,
            language=self.language,
            img_url=self.img_url
            )
        new.save()

    def manage(self):
        if Manager.is_new(self):
            Manager.write_bd(self)
            return
        else:
            Manager.update(self)
            return


class FlowListener():
    '''Parse soup and in case of new url ask Manager class
    to check if it is new one'''

    def __init__(self,
                url=config_url,
                # soup=None,
                log='tr_helper/last_article_log.txt'):
        self.log = log
        self.url = url
        self.last = open('tr_helper/last_article_log.txt').readlines()[-1].strip()
        # self.soup = soup or BeautifulSoup(requests.get(url).text, 'lxml')

    def start(self):
        self.soup = BeautifulSoup(requests.get(self.url).text, 'lxml')
        links = self.soup.table.find_all('a')
        urls = []
        # test_tracker = []
        for link in links:
            urls.append(link.get('href').strip())

        try:
            num = urls.index(self.last)
            if num:
                url = urls[:num][-1]
                print(url)
                Manager(url=url).manage()
                with open(self.log, 'a') as log:
                    log.write(url+'\n')

        except ValueError:
            pass

# def create_user():
#     user = User.objects.create_user(username, email, password)
#     user.is_translator = True
#     user.save()
#
#     cloud_account = User.objects.create_user(username, passowrd)
#     set_password
#     cloud_account.save()
#     account = CloudAccount(user=user, account=cloud_account, folder_name='some_name')
#     account.save()


urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

class PyMailCloudError(Exception):
    pass

    class AuthError(Exception):
        def __init__(self, message="Login or password is incorrect"):
            super(PyMailCloudError.AuthError, self).__init__(message)

    class NetworkError(Exception):
        def __init__(self, message="Connection failed"):
            super(PyMailCloudError.NetworkError, self).__init__(message)

    # class UnknownError(Exception):
    #     def __init__(self, message="WTF is going on?"):
    #         super(PyMailCloudError.UnknownError, self).__init__(message)

    class NotImplementedError(Exception):
        def __init__(self, message="The developer wants to sleep"):
            super(PyMailCloudError.NotImplementedError, self).__init__(message)
    class FileSizeError(Exception):
        def __init__(self, message="The file is bigger than 2 GB"):
            super(PyMailCloudError.FileSizeError, self).__init__(message)


class PyMailCloud:
    def __init__(self, login, password):

        self.session = requests.Session()
        self.login = login
        self.password = password
        self.downloadSource = None
        self.uploadTarget = None
        self.__recreate_token()

    def __get_download_source(self):
        dispatcher = self.session.get('https://cloud.mail.ru/api/v2/dispatcher',
                                      params={
                                          "token": self.token
                                      }, )
        if dispatcher.status_code is not 200:
            raise PyMailCloudError.NetworkError
        print(json.loads(dispatcher.content.decode("utf-8")))
        self.downloadSource = json.loads(dispatcher.content.decode("utf-8"))['body']['get'][0]['url']
        self.uploadTarget = json.loads(dispatcher.content.decode("utf-8"))['body']['upload'][0]['url']
        print('Acquired CDN Node')

    def __recreate_token(self):
        loginResponse = self.session.post("https://auth.mail.ru/cgi-bin/auth",
                                          data={
                                              "page": "http://cloud.mail.ru/",
                                              "Login": self.login,
                                              "Password": self.password
                                          },verify=False
                                          )

        if loginResponse.status_code == requests.codes.ok and loginResponse.history:
            getTokenResponse = self.session.post("https://cloud.mail.ru/api/v2/tokens/csrf")
            if getTokenResponse.status_code is not 200:
                raise PyMailCloudError.UnknownError
            self.token = json.loads(getTokenResponse.content.decode("utf-8"))['body']['token']
            print('Login successful')
            self.__get_download_source()
        else:
            raise PyMailCloudError.NetworkError()


    def upload_callback(self, monitor, progress):
        progress.total = monitor.len
        progress.update(8192)
        pass

    def upload_files(self, file, file_name, folder_name):
        progress = tqdm(unit='B')
        progress.desc = file_name

        destination = '/' + folder_name + '/' + file_name
        # if os.path.getsize(file['filename']) > 1024 * 1024 * 1024 * 2:
        #     raise PyMailCloudError.FileSizeError
        monitor = MultipartEncoderMonitor.from_fields(
            fields={'file': ('filename', file, 'application/octet-stream')},
            callback=lambda monitor: self.upload_callback(monitor, progress))
        upload_response = self.session.post(self.uploadTarget, data=monitor,
                          headers={'Content-Type': monitor.content_type},verify=False)
        if upload_response.status_code is not 200:
            raise PyMailCloudError.NetworkError

        print(monitor)
        hash, filesize = upload_response.content.decode("utf-8").split(';')[0], upload_response.content.decode("utf-8").split(';')[1][:-2]
        response = self.session.post("https://cloud.mail.ru/api/v2/file/add",  # "http://httpbin.org/post",
                                     data={
                                         "token": self.token,
                                         "home": destination,
                                         "conflict": 'rename',
                                         "hash": hash,
                                         "size": filesize,
                                     })
        return json.dumps(response.json(), sort_keys=True, indent=3, ensure_ascii=False)



def encode_docx(binary_text):
    docx = Document(BytesIO(binary_text))
    text = [paragraph.text for paragraph in docx.paragraphs]
    return text

def count_ammount_in_loaded(loaded_file):
    suffix = loaded_file.name.split('.')[-1]
    encode_by_suffix = {
        # 'doc' : encode_doc_txt,
        # 'txt' : encode_txt,
        'docx': encode_docx,
    }
    try:
        text = encode_by_suffix[suffix](loaded_file.read())
    except KeyError:
        return 'Unsupported format of uploaded file: can not count symbols_ammount'
    pure_text = [line for line in text if '://anf' not in line]
    ammount = sum(map(len, text))
    return ammount


def parse(url):
    parsed = NewArticle(url=url)
    lines = (parsed.title, parsed.lead, '', parsed.text, '', parsed.url)
    filename = parsed.url_title
    document = Document()
    for line in lines:
        document.add_paragraph(line)

    file = BytesIO()
    document.save(file)
    length = file.tell()
    file.seek(0)
    response = HttpResponse(
        file.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    response['Content-Disposition'] = 'attachment; filename={}.docx'.format(filename)
    response['Content-Length'] = length
    return response
